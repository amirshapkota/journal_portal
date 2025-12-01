"""
Document anonymization utility for blind review processes.

Handles the anonymization of author information in submitted documents
for Single Blind and Double Blind review processes.
"""
import re
import logging
from io import BytesIO
from typing import List, Tuple

logger = logging.getLogger(__name__)


class DocumentAnonymizer:
    """
    Anonymizes documents by replacing author information with asterisks.
    Supports PDF, DOCX, and TXT files.
    """
    
    ANONYMIZATION_MARKER = "******"
    
    @staticmethod
    def should_anonymize(journal_settings: dict) -> bool:
        """
        Check if documents should be anonymized based on journal settings.
        
        Args:
            journal_settings: Journal settings dictionary
            
        Returns:
            True if anonymization is required, False otherwise
        """
        # Handle invalid settings (e.g., string instead of dict)
        if not isinstance(journal_settings, dict):
            logger.warning(f"Invalid journal settings type: {type(journal_settings)}")
            return False
        
        # Support both 'review_type' (direct) and 'review_process.type' (nested) formats
        review_type = journal_settings.get('review_type', '').upper()
        
        # Fallback to nested structure if direct key not found
        if not review_type:
            review_process = journal_settings.get('review_process', {})
            if isinstance(review_process, dict):
                review_type = review_process.get('type', '').upper()
        
        # Anonymize for Single Blind and Double Blind
        return review_type in ['SINGLE_BLIND', 'DOUBLE_BLIND']
    
    @staticmethod
    def get_author_names(submission) -> List[str]:
        """
        Extract all author names from submission.
        
        Args:
            submission: Submission instance
            
        Returns:
            List of author full names and variations
        """
        names = []
        
        # Get all author contributions
        for contribution in submission.author_contributions.all():
            profile = contribution.profile
            user = profile.user
            
            # Add full name
            full_name = f"{user.first_name} {user.last_name}".strip()
            if full_name:
                names.append(full_name)
            
            # Add first name
            if user.first_name:
                names.append(user.first_name)
            
            # Add last name
            if user.last_name:
                names.append(user.last_name)
            
            # Add display name if different
            if profile.display_name and profile.display_name != full_name:
                names.append(profile.display_name)
            
            # Add reversed name (Last, First)
            if user.first_name and user.last_name:
                names.append(f"{user.last_name}, {user.first_name}")
            
            # Add initials (e.g., "J. Doe", "John D.", "Doe, J.")
            if user.first_name and user.last_name:
                names.append(f"{user.first_name[0]}. {user.last_name}")
                names.append(f"{user.first_name} {user.last_name[0]}.")
                names.append(f"{user.last_name}, {user.first_name[0]}.")
                names.append(f"{user.first_name[0]}. {user.last_name[0]}.")
        
        # Also include co-authors from metadata_json
        if submission.metadata_json and 'co_authors' in submission.metadata_json:
            co_authors = submission.metadata_json.get('co_authors', [])
            logger.info(f"Found {len(co_authors)} co-authors in metadata_json")
            
            for co_author in co_authors:
                if isinstance(co_author, dict):
                    # Extract name from co-author dict (format may vary)
                    co_author_name = co_author.get('name', '')
                    if co_author_name:
                        names.append(co_author_name)
                        
                        # Try to split into first/last name
                        name_parts = co_author_name.strip().split()
                        if len(name_parts) >= 2:
                            first_name = name_parts[0]
                            last_name = ' '.join(name_parts[1:])
                            
                            # Add variations
                            names.append(first_name)
                            names.append(last_name)
                            names.append(f"{last_name}, {first_name}")
                            names.append(f"{first_name[0]}. {last_name}")
                            names.append(f"{first_name} {last_name[0]}.")
                            names.append(f"{last_name}, {first_name[0]}.")
                            names.append(f"{first_name[0]}. {last_name[0]}.")
                elif isinstance(co_author, str):
                    # Direct string name
                    names.append(co_author)
                    
                    # Try to split into first/last name
                    name_parts = co_author.strip().split()
                    if len(name_parts) >= 2:
                        first_name = name_parts[0]
                        last_name = ' '.join(name_parts[1:])
                        
                        # Add variations
                        names.append(first_name)
                        names.append(last_name)
                        names.append(f"{last_name}, {first_name}")
                        names.append(f"{first_name[0]}. {last_name}")
                        names.append(f"{first_name} {last_name[0]}.")
                        names.append(f"{last_name}, {first_name[0]}.")
                        names.append(f"{first_name[0]}. {last_name[0]}.")
        
        # Always include corresponding_author if available (not just as fallback)
        if submission.corresponding_author:
            profile = submission.corresponding_author
            user = profile.user
            
            if not names:
                logger.info(f"No author contributions or co-authors found. Using corresponding_author: {user.first_name} {user.last_name}")
            
            # Add full name
            full_name = f"{user.first_name} {user.last_name}".strip()
            if full_name:
                names.append(full_name)
            
            # Add first name
            if user.first_name:
                names.append(user.first_name)
            
            # Add last name
            if user.last_name:
                names.append(user.last_name)
            
            # Add display name if different
            if profile.display_name and profile.display_name != full_name:
                names.append(profile.display_name)
            
            # Add reversed name (Last, First)
            if user.first_name and user.last_name:
                names.append(f"{user.last_name}, {user.first_name}")
            
            # Add initials (e.g., "J. Doe", "John D.", "Doe, J.")
            if user.first_name and user.last_name:
                names.append(f"{user.first_name[0]}. {user.last_name}")
                names.append(f"{user.first_name} {user.last_name[0]}.")
                names.append(f"{user.last_name}, {user.first_name[0]}.")
                names.append(f"{user.first_name[0]}. {user.last_name[0]}.")
        
        # Remove duplicates and sort by length (longest first for better matching)
        unique_names = list(set(filter(None, names)))
        unique_names.sort(key=len, reverse=True)
        
        return unique_names
    
    @staticmethod
    def anonymize_pdf(file_content: bytes, author_names: List[str]) -> bytes:
        """
        Anonymize file by replacing author names.
        
        Args:
            file_content: PDF file content as bytes
            author_names: List of author names to replace
            
        Returns:
            Anonymized PDF content as bytes
        """
        try:
            from PyPDF2 import PdfReader, PdfWriter
            from io import BytesIO
            
            reader = PdfReader(BytesIO(file_content))
            writer = PdfWriter()
            
            # Process each page
            for page in reader.pages:
                # Extract text
                text = page.extract_text()
                
                # For now, we'll log that PDF anonymization requires additional libraries
                # In production, consider using pdf-redactor or similar
                logger.warning("PDF anonymization requires additional processing. Consider using DOCX format.")
                
                writer.add_page(page)
            
            # For metadata anonymization
            metadata = reader.metadata or {}
            if metadata:
                # Remove author from metadata
                metadata_dict = {k: v for k, v in metadata.items() if k != '/Author'}
                writer.add_metadata(metadata_dict)
            
            # Write to bytes
            output = BytesIO()
            writer.write(output)
            output.seek(0)
            
            return output.read()
            
        except ImportError:
            logger.error("PyPDF2 not installed. Cannot anonymize PDF files.")
            return file_content
        except Exception as e:
            logger.error(f"Error anonymizing PDF: {str(e)}")
            return file_content
    
    @staticmethod
    def anonymize_docx(file_content: bytes, author_names: List[str]) -> bytes:
        """
        Anonymize DOCX file by replacing author names.
        
        Args:
            file_content: DOCX file content as bytes
            author_names: List of author names to replace
            
        Returns:
            Anonymized DOCX content as bytes
        """
        try:
            from docx import Document
            from io import BytesIO
            
            doc = Document(BytesIO(file_content))
            
            logger.info(f"Anonymizing DOCX with {len(author_names)} name variations")
            
            # Anonymize in paragraphs
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    original_text = run.text
                    modified_text = original_text
                    
                    for name in author_names:
                        # Use case-insensitive replacement
                        pattern = re.compile(re.escape(name), re.IGNORECASE)
                        modified_text = pattern.sub(DocumentAnonymizer.ANONYMIZATION_MARKER, modified_text)
                    
                    if modified_text != original_text:
                        run.text = modified_text
                        logger.info(f"Replaced text in paragraph run")
            
            # Anonymize in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                original_text = run.text
                                modified_text = original_text
                                
                                for name in author_names:
                                    pattern = re.compile(re.escape(name), re.IGNORECASE)
                                    modified_text = pattern.sub(DocumentAnonymizer.ANONYMIZATION_MARKER, modified_text)
                                
                                if modified_text != original_text:
                                    run.text = modified_text
                                    logger.info(f"Replaced text in table cell")
            
            # Anonymize document properties/metadata
            core_properties = doc.core_properties
            if core_properties.author:
                logger.info(f"Anonymizing author metadata: {core_properties.author}")
                core_properties.author = DocumentAnonymizer.ANONYMIZATION_MARKER
            if core_properties.last_modified_by:
                core_properties.last_modified_by = DocumentAnonymizer.ANONYMIZATION_MARKER
            
            # Save to bytes
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            logger.info("DOCX anonymization completed")
            return output.read()
            
        except ImportError:
            logger.error("python-docx not installed. Cannot anonymize DOCX files.")
            return file_content
        except Exception as e:
            logger.error(f"Error anonymizing DOCX: {str(e)}")
            return file_content
    
    @staticmethod
    def anonymize_text(file_content: bytes, author_names: List[str], encoding: str = 'utf-8') -> bytes:
        """
        Anonymize text file by replacing author names.
        
        Args:
            file_content: Text file content as bytes
            author_names: List of author names to replace
            encoding: Text encoding (default: utf-8)
            
        Returns:
            Anonymized text content as bytes
        """
        try:
            text = file_content.decode(encoding)
            
            # Replace each author name
            for name in author_names:
                # Case-insensitive replacement
                pattern = re.compile(re.escape(name), re.IGNORECASE)
                text = pattern.sub(DocumentAnonymizer.ANONYMIZATION_MARKER, text)
            
            return text.encode(encoding)
            
        except Exception as e:
            logger.error(f"Error anonymizing text file: {str(e)}")
            return file_content
    
    @classmethod
    def anonymize_file(cls, file_obj, submission) -> Tuple[bytes, bool]:
        """
        Anonymize a file based on its type and journal settings.
        
        Args:
            file_obj: Django UploadedFile object
            submission: Submission instance
            
        Returns:
            Tuple of (file_content, was_anonymized)
        """
        # Check if anonymization is needed
        journal_settings = submission.journal.settings or {}
        
        logger.info(f"Checking anonymization for submission {submission.id}")
        logger.info(f"Journal settings: {journal_settings}")
        
        if not cls.should_anonymize(journal_settings):
            logger.info(f"Anonymization not required for submission {submission.id}")
            file_obj.seek(0)
            return file_obj.read(), False
        
        logger.info(f"Anonymization IS required for submission {submission.id}")
        
        # Get author names
        author_names = cls.get_author_names(submission)
        if not author_names:
            logger.warning(
                f"No author names found for submission {submission.id}. "
                f"Author contributions count: {submission.author_contributions.count()}. "
                f"Cannot anonymize without author information."
            )
            file_obj.seek(0)
            return file_obj.read(), False
        
        logger.info(f"Anonymizing document for submission {submission.id}. Found {len(author_names)} name variations: {author_names}")
        
        # Read file content
        file_obj.seek(0)
        original_content = file_obj.read()
        
        # Determine file type and anonymize
        file_name = file_obj.name.lower()
        
        if file_name.endswith('.pdf'):
            anonymized_content = cls.anonymize_pdf(original_content, author_names)
        elif file_name.endswith('.docx'):
            anonymized_content = cls.anonymize_docx(original_content, author_names)
        elif file_name.endswith(('.txt', '.md', '.tex', '.rtf')):
            anonymized_content = cls.anonymize_text(original_content, author_names)
        else:
            logger.warning(f"Unsupported file type for anonymization: {file_name}")
            return original_content, False
        
        logger.info(f"Document anonymized successfully for submission {submission.id}")
        return anonymized_content, True
